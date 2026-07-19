"""Trích text theo TRANG + phân loại trang: text | scan | hybrid (theo mật độ ký tự).

- .md/.txt: đọc thẳng, 1 "trang".
- .pdf: PyMuPDF; trang scan/hybrid => Vision OCR (render trang thành PNG, gửi VLM
  đọc thành Markdown — giữ bảng, dấu tiếng Việt, bỏ mộc/chữ ký).
- .docx: python-docx — đoạn văn + BẢNG (chuyển Markdown) theo đúng thứ tự trong file.
  Word không có "trang" thật (trang chỉ xác định khi render) => gom thành trang LOGIC:
  ngắt ở Heading 1/2 hoặc ~3000 ký tự, để citation "(nguồn: x.docx, trang n)" vẫn có nghĩa.
- .doc (Word 97-2003, nhị phân): KHÔNG hỗ trợ — mở Word "Save As .docx" rồi nạp.
- Tài liệu confidential: KHÔNG gửi ảnh trang lên cloud để OCR -> trang scan bị BỎ QUA
  (đã bỏ Vision local; muốn OCR tài liệu mật thì phải có kênh OCR nội bộ riêng).
"""
from dataclasses import dataclass
from pathlib import Path

import config
from rag.text.vi import normalize

_OCR_PROMPT = (
    "Đây là ảnh một trang tài liệu tiếng Việt (có thể là văn bản pháp luật, hợp đồng...). "
    "Hãy chép lại TOÀN BỘ nội dung thành Markdown: giữ nguyên dấu tiếng Việt, "
    "giữ cấu trúc bảng bằng bảng Markdown, giữ tiêu đề/điều/khoản. "
    "BỎ QUA con dấu, chữ ký, watermark. Chỉ trả về nội dung, không giải thích."
)


@dataclass
class Page:
    number: int   # 1-based
    text: str
    kind: str     # text | scan | hybrid | error
    status: str = "ok"   # ok | ocr_failed | ocr_skipped | blank | error
    note: str = ""       # lý do khi status != ok (để ghi báo cáo per-trang)


def extract_pages(path: str, confidential: bool = False) -> list[Page]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        text = normalize(p.read_text(encoding="utf-8", errors="replace"))
        if text.strip():
            return [Page(1, text, "text")]
        return [Page(1, "", "text", "blank", "file rỗng")]
    if suffix == ".pdf":
        return _extract_pdf(p, confidential)
    if suffix == ".docx":
        return _extract_docx(p)
    if suffix == ".doc":
        raise ValueError(
            "Định dạng .doc (Word 97-2003) chưa hỗ trợ — mở bằng Word, "
            "Save As sang .docx rồi nạp lại."
        )
    raise ValueError(f"Định dạng chưa hỗ trợ: {suffix} (hỗ trợ .pdf/.docx/.txt/.md)")


def _extract_pdf(p: Path, confidential: bool) -> list[Page]:
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError(
            "Thiếu PyMuPDF để đọc PDF — chạy: pip install PyMuPDF"
        ) from e

    pages: list[Page] = []
    with fitz.open(p) as doc:
        total = doc.page_count
        for i, page in enumerate(doc, start=1):
            # Một trang hỏng KHÔNG được giết cả tài liệu — bắt lỗi, ghi nhận, đi tiếp.
            try:
                text = normalize(page.get_text().strip())
            except Exception as e:  # trang PDF lỗi/corrupt
                pages.append(Page(i, "", "error", "error", str(e)[:200]))
                print(f"[extract] trang {i}/{total}: LỖI đọc — {str(e)[:150]}")
                continue
            density = len(text)
            if density >= config.PAGE_TEXT_MIN_CHARS:
                pages.append(Page(i, text, "text"))
                continue
            kind = "scan" if density == 0 else "hybrid"
            ocr, st, note = _ocr_page(page, confidential)
            merged = normalize((text + "\n\n" + ocr).strip() if kind == "hybrid" else ocr)
            if merged:  # có nội dung dùng được (OCR ok, hoặc hybrid còn phần text gốc)
                pages.append(Page(i, merged, kind, "ok", note if st != "ok" else ""))
                print(f"[ocr] trang {i}/{total} ({kind}): OK ({len(merged)} ký tự)")
            else:  # không trích được gì -> ghi đúng lý do (skipped/failed/blank)
                pages.append(Page(i, "", kind, st, note))
                print(f"[ocr] trang {i}/{total} ({kind}): {st.upper()} — {note}")
    return pages


DOCX_PAGE_CHARS = 3000  # ngưỡng gom trang logic cho .docx


def _extract_docx(p: Path) -> list[Page]:
    try:
        import docx
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as e:
        raise RuntimeError(
            "Thiếu python-docx để đọc .docx — chạy: pip install python-docx"
        ) from e

    d = docx.Document(str(p))
    if hasattr(d, "iter_inner_content"):  # python-docx >= 1.1: đúng thứ tự xen kẽ
        items = list(d.iter_inner_content())
    else:  # bản cũ: đoạn văn trước, bảng sau (mất thứ tự xen kẽ nhưng không mất nội dung)
        items = list(d.paragraphs) + list(d.tables)

    blocks: list[tuple[bool, str]] = []  # (is_big_heading, text)
    for item in items:
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            sid = (item.style.style_id or "") if item.style is not None else ""
            # style_id không bị localize như style.name -> nhận diện Heading ổn định
            is_heading = sid.startswith("Heading") and sid[-1:] in ("1", "2")
            blocks.append((is_heading, text))
        elif isinstance(item, Table):
            md = _table_to_markdown(item)
            if md:
                blocks.append((False, md))

    pages: list[Page] = []
    cur: list[str] = []
    cur_len = 0
    for is_heading, text in blocks:
        if cur and (is_heading or cur_len + len(text) > DOCX_PAGE_CHARS):
            pages.append(Page(len(pages) + 1, normalize("\n\n".join(cur)), "text"))
            cur, cur_len = [], 0
        cur.append(text)
        cur_len += len(text)
    if cur:
        pages.append(Page(len(pages) + 1, normalize("\n\n".join(cur)), "text"))
    return pages


def _table_to_markdown(table) -> str:
    """Bảng Word -> bảng Markdown (giữ được cấu trúc khi chunk + đưa vào context)."""
    rows: list[list[str]] = []
    for r in table.rows:
        cells = [" ".join(c.text.split()).replace("|", "\\|") for c in r.cells]
        rows.append(cells)
    rows = [r for r in rows if any(cell for cell in r)]
    if not rows:
        return ""
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    lines = ["| " + " | ".join(rows[0]) + " |", "|" + " --- |" * ncol]
    lines += ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return "\n".join(lines)


def _ocr_page(page, confidential: bool) -> tuple[str, str, str]:
    """Render trang -> PNG -> Vision OCR. Trả về (text, status, note):
    - ("...", "ok", "")            OCR đọc được nội dung
    - ("", "ocr_skipped", note)    Vision offline -> trang scan bị bỏ qua (không phải lỗi)
    - ("", "ocr_failed", note)     render/gọi VLM ném lỗi
    - ("", "blank", note)          VLM trả rỗng (trang trắng / ảnh không đọc được)
    """
    from rag.generate.llm import vision_ocr

    provider = config.VISION_PROVIDER
    if confidential:
        # tài liệu MẬT: cấm gửi ảnh trang lên cloud để OCR (đã bỏ Vision local) -> bỏ qua
        return "", "ocr_skipped", "tài liệu mật: không gửi ảnh lên cloud để OCR (trang scan bị bỏ)"
    try:
        png = page.get_pixmap(dpi=150).tobytes("png")
        text = vision_ocr(png, _OCR_PROMPT, provider=provider)
    except Exception as e:
        return "", "ocr_failed", str(e)[:200]
    text = (text or "").strip()
    if not text:
        return "", "blank", "OCR không trích được nội dung (trang trắng hoặc ảnh mờ)"
    return text, "ok", ""
